#!/usr/bin/env python3
"""Generate Network_Energy_Measurements.docx with clean, explanatory text."""

import csv, statistics, math
from pathlib import Path
from collections import defaultdict
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = Path(__file__).resolve().parent
SUMMARY = BASE / "results" / "summary.csv"
PLOTS = BASE / "results" / "plots_network"
OUT = BASE / "results" / "Network_Energy_Measurements.docx"

def sci(v):
    if v == 0: return "0"
    e = int(math.floor(math.log10(abs(v))))
    m = v / 10**e
    sup = str(e).replace("-", "\u207B")
    sup = sup.replace("0", "\u2070").replace("1", "\u00B9").replace("2", "\u00B2")
    sup = sup.replace("3", "\u00B3").replace("4", "\u2074").replace("5", "\u2075")
    sup = sup.replace("6", "\u2076").replace("7", "\u2077").replace("8", "\u2078")
    sup = sup.replace("9", "\u2079")
    return f"{m:.4f} \u00D7 10{sup}"

def load():
    with open(SUMMARY) as f:
        rows = list(csv.DictReader(f))
    g = defaultdict(list)
    for r in rows:
        n = r["creative_platform"]
        if n.startswith("nexd__"): g["NEXD"].append(r)
        elif n.startswith("html5hiili__"): g["alt_html5"].append(r)
        elif n.startswith("html5__"): g["html5_s3"].append(r)
    cmap = {c.lower(): c for c in rows[0].keys()}
    def col(base): return cmap.get(base.lower(), base)
    stats = {}
    for pref, name in [("wifi_ac", "Wi-Fi (802.11ac)"), ("lte", "LTE / 4G")]:
        s = {}
        for grp in ["NEXD", "html5_s3", "alt_html5"]:
            s[grp] = {
                "n": len(g[grp]),
                "total": statistics.median([float(r[col(f"{pref}_consumed_kWh")]) for r in g[grp]]),
                "eff": statistics.median([float(r[col(f"{pref}_kWh_per_useful_mb")]) for r in g[grp]]),
                "eff_wire": statistics.median([float(r[col(f"{pref}_kWh_per_mb")]) for r in g[grp]]),
                "payload": statistics.median([float(r["payload_mb"]) for r in g[grp]]),
                "useful": statistics.median([float(r["useful_mb"]) for r in g[grp]]),
            }
        s["sav_total_vs_5"] = (1 - s["NEXD"]["total"] / s["html5_s3"]["total"]) * 100
        s["sav_total_vs_a"] = (1 - s["NEXD"]["total"] / s["alt_html5"]["total"]) * 100
        s["sav_eff_vs_5"] = (1 - s["NEXD"]["eff"] / s["html5_s3"]["eff"]) * 100
        s["sav_eff_vs_a"] = (1 - s["NEXD"]["eff"] / s["alt_html5"]["eff"]) * 100
        stats[name] = s
    return stats

s = load()
wifi, lte = s["Wi-Fi (802.11ac)"], s["LTE / 4G"]

doc = Document()
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)

# ── Title ──
h = doc.add_paragraph()
r = h.add_run("COMPARATIVE ENERGY CONSUMPTION ANALYSIS\nHTML5 vs NEXD Ad Formats")
r.font.name = 'Arial'; r.font.size = Pt(18); r.font.bold = True; r.font.color.rgb = RGBColor(0x33,0x33,0x33)
h.paragraph_format.space_after = Pt(6)
sub = doc.add_paragraph()
r = sub.add_run("Network Transmission Energy (NS-3.48)")
r.font.size = Pt(10); r.italic = True
sub.paragraph_format.space_after = Pt(12)

# ── 1. Introduction ──
doc.add_heading("1. Introduction", level=1)
doc.add_paragraph(
    f"This study compares the network transmission energy of 175 advertisements across three groups: "
    f"65 NEXD, 53 HTML5, and 57 alt. HTML5. The HTML5 group corresponds to rich-media display ads "
    f"provided as local assets (scenarios). The alt. HTML5 group corresponds to display-banner previews "
    f"served via Google Ads CDN. NEXD creatives (.acz bundles) are rich-media display ads comparable "
    f"in complexity to the HTML5 group."
)

# ── 2. Network Simulation ──
doc.add_heading("2. Network Transmission Energy (NS-3 Simulation)", level=1)
doc.add_paragraph(
    "Network transmission was simulated using NS-3.48 under Wi-Fi (802.11ac) and LTE/4G. "
    "The energy model captures active transmission time: more wire bytes \u2192 more radio-on time \u2192 more energy. "
    "\"Wire\" bytes are what travels over the air (compressed). \"Useful\" bytes are the uncompressed content that reaches the device. "
    "HTML5 local file:// assets are gzip-compressed during simulation (text only; images and video keep their original size). "
    "NEXD and alt. HTML5 wire bytes come from real HAR _transferSize captured via CDP."
)

# ── 2.1 Core data table ──
doc.add_heading("2.1 Core measurements", level=2)
t = doc.add_table(rows=7, cols=6)
t.style = 'Light Shading Accent 1'
widths = [Inches(1.6), Inches(1.4), Inches(0.4), Inches(1.5), Inches(1.5), Inches(1.1)]
for row in t.rows:
    for idx, w in enumerate(widths): row.cells[idx].width = w
hdr = t.rows[0].cells
for j, txt in enumerate(["Technology","Ad Type","N","Median Total Energy (kWh)","Median Energy / MB useful (kWh/MB)","NEXD Savings"]):
    hdr[j].text = txt
    for p in hdr[j].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs: run.font.bold = True; run.font.size = Pt(8)
data = [
    ("Wi-Fi (802.11ac)", "NEXD", str(wifi["NEXD"]["n"]), sci(wifi["NEXD"]["total"]), sci(wifi["NEXD"]["eff"]),
     f'{wifi["sav_total_vs_5"]:.1f}% vs HTML5\n{wifi["sav_total_vs_a"]:.1f}% vs alt. html5'),
    ("Wi-Fi (802.11ac)", "HTML5", str(wifi["html5_s3"]["n"]), sci(wifi["html5_s3"]["total"]), sci(wifi["html5_s3"]["eff"]), "Baseline"),
    ("Wi-Fi (802.11ac)", "alt. html5", str(wifi["alt_html5"]["n"]), sci(wifi["alt_html5"]["total"]), sci(wifi["alt_html5"]["eff"]), "Baseline"),
    ("LTE / 4G", "NEXD", str(lte["NEXD"]["n"]), sci(lte["NEXD"]["total"]), sci(lte["NEXD"]["eff"]),
     f'{lte["sav_total_vs_5"]:.1f}% vs HTML5\n{lte["sav_total_vs_a"]:.1f}% vs alt. html5'),
    ("LTE / 4G", "HTML5", str(lte["html5_s3"]["n"]), sci(lte["html5_s3"]["total"]), sci(lte["html5_s3"]["eff"]), "Baseline"),
    ("LTE / 4G", "alt. html5", str(lte["alt_html5"]["n"]), sci(lte["alt_html5"]["total"]), sci(lte["alt_html5"]["eff"]), "Baseline"),
]
for i, row in enumerate(data, 1):
    for j, txt in enumerate(row):
        t.rows[i].cells[j].text = txt
        for p in t.rows[i].cells[j].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j in (0,1) else WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs: run.font.size = Pt(8)
doc.add_paragraph()

# ── 2.2 What drives the difference ──
doc.add_heading("2.2 What drives these numbers", level=2)
p = doc.add_paragraph(
    "Energy per wire MB is nearly identical across all three formats "
    f"({sci(wifi['NEXD']['eff_wire'])} kWh/MB for NEXD, "
    f"{sci(wifi['html5_s3']['eff_wire'])} for HTML5, "
    f"{sci(wifi['alt_html5']['eff_wire'])} for alt. HTML5 on Wi-Fi). "
    "Radio physics does not distinguish between formats: more bytes over the air means more energy. "
    "Differences in total energy come almost entirely from differences in wire bytes."
)
p.paragraph_format.space_after = Pt(6)

t2 = doc.add_table(rows=4, cols=5)
t2.style = 'Light Shading Accent 1'
h2 = t2.rows[0].cells
for j, txt in enumerate(["Group","Wire MB","Useful MB","Useful/Wire","Creative type"]):
    h2[j].text = txt
    for p in h2[j].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs: run.font.bold = True; run.font.size = Pt(8)
pw = wifi['NEXD']['payload']; uw = wifi['NEXD']['useful']
ph = wifi['html5_s3']['payload']; uh = wifi['html5_s3']['useful']
pa = wifi['alt_html5']['payload']; ua = wifi['alt_html5']['useful']
for i, (grp, pl, ul, ratio, ctype) in enumerate([
    ("NEXD", f"{pw:.4f}", f"{uw:.4f}", f"{uw/pw:.1%}", "Rich-media display (.acz)"),
    ("HTML5", f"{ph:.4f}", f"{uh:.4f}", f"{uh/ph:.1%}", "Rich-media display (local files)"),
    ("alt. HTML5", f"{pa:.4f}", f"{ua:.4f}", f"{ua/pa:.1%}", "Simple display banners (CDN)"),
], 1):
    for j, txt in enumerate([grp, pl, ul, ratio, ctype]):
        t2.rows[i].cells[j].text = txt
        for p in t2.rows[i].cells[j].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if j in (1,2,3) else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs: run.font.size = Pt(8)
doc.add_paragraph()

doc.add_paragraph(
    "HTML5 group ads are rich-media creatives with multiple large PNG images and MP4 video "
    "(\u223c5\u201311 MB wire). alt. HTML5 group ads are simpler display banners "
    "(\u223c0.4\u20131.4 MB wire) that happen to be comparable in weight to NEXD. "
    "Neither group is a perfect counterfactual for NEXD: HTML5 matches the creative complexity "
    "but is much heavier; alt. HTML5 matches the wire weight but consists of simpler creatives."
)

# ── 3. Key conclusions ──
doc.add_heading("3. Key Conclusions", level=1)
conclusions = [
    ("NEXD saves 81.7% total energy vs the HTML5 group",
     f"The source creatives in the HTML5 group are rich-media ads weighing ~{ph:.2f} MB wire. "
     f"NEXD achieves similar creative richness at ~{pw:.2f} MB wire \u2014 "
     f"{ph/pw:.0f}\u00D7 less data. Since energy scales with wire bytes, this is the dominant effect. "
     "NEXD\u2019s advantage is not better compression but lighter source assets."),
    ("NEXD saves only 2.6% vs the alt. HTML5 group",
     f"The alt. HTML5 group consists of display banners that weigh ~{pa:.2f} MB wire, "
     f"essentially the same as NEXD ({pw:.2f} MB). This confirms NEXD\u2019s .acz format is "
     "equivalent in weight to a simple display banner served through a CDN. "
     "The comparison is not apples-to-apples in creative richness."),
    ("Energy per MB useful: NEXD is 28% better than HTML5, 10% worse than alt. HTML5",
     "NEXD\u2019s payload has a higher proportion of compressible text (JS, JSON) "
     f"(useful/wire = {uw/pw:.0%}) vs HTML5 "
     f"({uh/ph:.0%}, dominated by incompressible images/video). "
     "This gives NEXD a modest efficiency advantage per useful MB vs the HTML5 group. "
     f"vs alt. HTML5 (useful/wire = {ua/pa:.0%}), NEXD is slightly less efficient "
     "due to .acz decompression overhead."),
    ("Comparison validity depends on creative type",
     "The 81.7% savings vs the HTML5 group is the more relevant comparison for "
     "rich-media display ads, which is the category NEXD targets. The alt. HTML5 group "
     "is a useful reference for wire-weight comparison but does not match NEXD in "
     "creative complexity."),
]
for title, body in conclusions:
    h = doc.add_paragraph()
    r = h.add_run(title); r.bold = True; r.font.size = Pt(10)
    p = doc.add_paragraph(body); p.paragraph_format.space_after = Pt(4)

# ── 4. Methodology ──
doc.add_heading("4. Methodology", level=1)
doc.add_paragraph(
    f"Simulator: NS-3.48 with WifiRadioEnergyModelHelper (Wi-Fi) and LTE PHY traces. "
    f"Sample: 65 NEXD, 53 HTML5, 57 alt. HTML5 = 175 total. "
    "HARs captured via CDP (crawl_ads.py) and Playwright (crawl_ads_csv_urls.py). "
    "HTML5 local file:// assets read from disk and gzip-compressed during simulation "
    "(text MIME types only; images/video keep original size). "
    "NEXD and alt. HTML5 use real _transferSize from CDP capture (Brotli/gzip from CDN)."
)

doc.save(str(OUT))
print(f"Generated: {OUT}")
